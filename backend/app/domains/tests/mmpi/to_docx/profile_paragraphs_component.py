from docx.document import Document

from app.schemas.test_result import TestResultDto
from app.utils.results_to_docx.document_component import DocumentComponent
from app.utils.results_to_docx.document_formatter import DocumentFormatter


class ProfileParagraphsComponent(DocumentComponent):
    """Renders paragraphs for profile types and inclinations for MMPI test"""

    def __init__(self, document: Document, formatter: DocumentFormatter, test_result: TestResultDto):
        super().__init__(document, formatter)
        self.profile_types = test_result.verdict.profile_types
        self.profile_inclinations = test_result.verdict.profile_inclinations

    def render(self) -> None:
        self._add_profiles_paragraph(self.profile_types, "Типи профілю")
        self._add_profiles_paragraph(self.profile_inclinations, "Нахил профілю")

    def _add_profiles_paragraph(self, types: list[str], title: str):
        if not types or len(types) == 0:
            return

        p_profile_types = self.doc.add_paragraph()
        p_profile_types.add_run(f"{title}: ").bold = True
        p_profile_types.add_run(", ".join(types) + ".")
        self.formatter.style_paragraph(p_profile_types)
