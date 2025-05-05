from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.schemas.test_result import TestResultDto
from app.schemas.user_auth import UserDto
from app.utils.results_to_docx.document_component import DocumentComponent
from app.utils.results_to_docx.document_formatter import DocumentFormatter


class HeaderComponent(DocumentComponent):
    """
    Handles the document header section.
    Shows test title, information about patient and test pass date
    """

    def __init__(self,
                 document: Document, formatter: DocumentFormatter,
                 patient: UserDto, test_result: TestResultDto):
        super().__init__(document, formatter)
        self.patient = patient
        self.test_result = test_result

    def render(self):
        self._add_title()
        self._add_patient_credentials()
        self._add_age()
        self._add_pass_date()

    def _add_title(self):
        header = self.doc.add_paragraph()
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        header.add_run(f"Бланк {self.test_result.test.name}").bold = True
        self.formatter.set_font_size(header, 14)

    def _add_patient_credentials(self):
        p = self.patient
        credentials = self.doc.add_paragraph(f"П.І.Б: {p.surname} "
                                             f"{p.name if p.name else ''} "
                                             f"{p.patronymic if p.patronymic else ''}")
        self.formatter.style_paragraph(credentials)

    def _add_age(self):
        (age, age_ending) = self.patient.get_age()
        age_paragraph = self.doc.add_paragraph(f"Вік пацієнта: {age} {age_ending}")
        self.formatter.style_paragraph(age_paragraph)

    def _add_pass_date(self):
        pass_date = self.test_result.passed_at.strftime("%d.%m.%Y %H:%M")
        pass_paragraph = self.doc.add_paragraph(f"Дата проходження: {pass_date}")
        self.formatter.style_paragraph(pass_paragraph)
