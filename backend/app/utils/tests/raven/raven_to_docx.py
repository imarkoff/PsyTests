from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.schemas.test.test_history_results import Answer, Results
from app.utils.results_to_docx import ResultsToDocx


class RavenToDocx(ResultsToDocx):
    def _fill_content(self):
        self._create_verdict()

        max_answers = self._calculate_max_answers(self.test_result.results)
        self.table = self._create_table(max_answers)
        self._fill_table()

    @staticmethod
    def _calculate_max_answers(results: Results):
        max_answers = 0

        for module, answers in results.items():
            if len(answers) > max_answers:
                max_answers = len(answers)

        return max_answers

    def _create_verdict(self):
        verdict = self.test_result.verdict.get("_") if self.test_result.verdict else None
        marks_unit = self.test_result.test.marks_unit
        if verdict:
            self.doc.add_paragraph(f"Висновок: {verdict} {marks_unit if marks_unit else ""}")

    def _create_table(self, max_answers: int):
        table = self.doc.add_table(rows=0, cols=max_answers + 2)
        table.style = "Table Grid"

        table.add_row()

        hdr_cells = table.rows[0].cells

        for i in range(max_answers):
            index_cell = hdr_cells[i + 1]
            index_cell.text = str(i + 1)
            index_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            index_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            self._set_font_size(index_cell, 11)

        total_cell = hdr_cells[-1]
        total_cell.text = "Бали"
        self._set_font_size(total_cell, 11)

        return table

    def _fill_table(self):
        correct_points = 0
        total_points = 0

        for module, answers in self.test_result.results.items():
            row_cells = self.table.add_row().cells
            self._fill_cell(row_cells[0], module)

            module_correct_points, module_total_points = self._fill_answers(row_cells, answers)

            self._fill_cell(row_cells[-1], module_correct_points, align=WD_ALIGN_PARAGRAPH.RIGHT)

            total_points += module_total_points
            correct_points += module_correct_points

        self._add_total_row(correct_points, total_points)

    def _fill_cell(self, cell, text, align=WD_ALIGN_PARAGRAPH.LEFT):
        cell.text = str(text)
        self._set_font_size(cell, 11)
        cell.paragraphs[0].alignment = align
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    def _fill_answers(self, row_cells, answers: list[Answer]):
        module_correct_points = 0
        module_total_points = 0

        for i, answer in enumerate(answers):
            user_answer = answer["user_answer"]
            correct_answer = answer["correct_answer"]
            points = answer["points"]

            str_user_answer = "" if user_answer is None else f"{user_answer + 1}"
            str_correct_answer = f" ({correct_answer + 1})" if user_answer is not None and user_answer != correct_answer else ""

            self._fill_cell(
                row_cells[i + 1],
                str_user_answer + str_correct_answer,
                align=WD_ALIGN_PARAGRAPH.CENTER
            )

            module_total_points += points
            if user_answer == correct_answer:
                module_correct_points += points

        return module_correct_points, module_total_points

    def _add_total_row(self, correct_points, total_points):
        row = self.table.add_row()
        results = row.cells

        results_cell = results[0]
        self._fill_cell(
            results_cell,
            f"Загальний бал: {correct_points} з {total_points}",
            align=WD_ALIGN_PARAGRAPH.RIGHT
        )
        results_cell.merge(results[-1])