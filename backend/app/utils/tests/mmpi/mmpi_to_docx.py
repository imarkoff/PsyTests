import os

from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL

from app.utils.results_to_docx import ResultsToDocx


class MMPIToDocx(ResultsToDocx):
    def _fill_content(self):
        chart_path = self._get_result_bars()
        self.doc.add_picture(chart_path, width=None, height=None)
        if os.path.exists(chart_path):
            os.remove(chart_path)

        self._get_raw_table()
        self._add_space_after_table()

        self._add_profiles_row(self.test_result.verdict.get("profile_types"), "Типи профілю")
        self._add_profiles_row(self.test_result.verdict.get("profile_inclinations"), "Нахил профілю")
        self._add_verdicts_table()

    def _get_result_bars(self):
        from matplotlib import pyplot as plt
        import tempfile

        plt.style.use('_mpl-gallery')

        (x, y) = zip(*self.test_result.verdict.get("converted").items())
        bar_colors = ['red' if val <= 30 or val >= 70 else 'green' for val in y]

        fig, ax = plt.subplots(figsize=(7.5, 3))
        bars = ax.bar(x, y, width=1, edgecolor="white", linewidth=1, color=bar_colors)
        ax.set(ylim=(0, 130))

        ax.bar_label(bars, fmt='%i', label_type="center")

        plt.tight_layout()
        plt.subplots_adjust(left=0.05, right=0.95, top=0.9, bottom=0.1)

        ax.grid(axis='y', alpha=0.6)
        ax.grid(axis='x', alpha=0)
        ax.set_title("Конвертовані результати тесту")

        temp_name = f'mmpi_plot_{self.test_result.id}.png'
        temp_file = os.path.join(tempfile.gettempdir(), temp_name)
        plt.savefig(temp_file, dpi=300, bbox_inches='tight')
        plt.close(fig)

        return temp_file


    def _get_raw_table(self):
        (scales, values) = zip(*self.test_result.verdict.get("raw").items())

        table = self.doc.add_table(rows=0, cols=len(scales))
        table.style = "Table Grid"

        header = table.add_row().cells
        header[0].text = "Сирі результати"
        header[0].merge(header[-1])

        scales_row = table.add_row().cells
        for i, scale in enumerate(scales):
            self._fill_cell(scales_row[i], scale, align=WD_ALIGN_PARAGRAPH.CENTER)

        values_row = table.add_row().cells
        for i, value in enumerate(values):
            self._fill_cell(values_row[i], value, align=WD_ALIGN_VERTICAL.CENTER)

        return table


    def _add_profiles_row(self, types: list[str], title: str):
        if not types or len(types) == 0:
            return

        p_profile_types = self.doc.add_paragraph()
        p_profile_types.add_run(f"{title}: ").bold = True
        p_profile_types.add_run(", ".join(types) + ".")
        self._style_paragraph(p_profile_types)

    def _add_verdicts_table(self):
        table = self.doc.add_table(rows=0, cols=2)
        table.style = "Table Grid"

        table.columns[0].width = Cm(1)
        table.columns[1].width = Cm(18)

        header = table.add_row().cells
        header[0].text = "Висновки"
        header[0].merge(header[-1])

        for scale, description in self.test_result.verdict.get("scale_verdicts").items():
            row = table.add_row().cells
            row[0].text = scale
            row[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            row[1].text = description

        return table

