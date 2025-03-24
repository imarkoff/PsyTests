import os.path
import tempfile
from typing import TYPE_CHECKING

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from app.schemas.user_auth import UserDto

if TYPE_CHECKING:
    from app.schemas.test_result import TestResultDto

temp_dir = tempfile.gettempdir()

class ResultsToDocx:
    """
    Base class for creating docx file with test results
    Logic for filling the document should be implemented in _fill_content method
    """

    path: str = None

    def __init__(self, test_result: 'TestResultDto', user_info: UserDto):
        self.test_result = test_result
        self.user_info = user_info

        self.doc = Document()

        self.doc.sections[0].left_margin = Pt(42)
        self.doc.sections[0].right_margin = Pt(32)

        self._create_header()
        self._fill_content()

        self.path = self._save()

    def __del__(self):
        if self.path and os.path.exists(self.path):
            os.remove(self.path)

    def _fill_content(self):
        """
        Method for filling the document with content.
        Should be implemented in child classes
        """
        pass

    def _save(self):
        """
        Save document to file and return path
        """
        surname = self.user_info.surname
        name = self.user_info.name
        patronymic = self.user_info.patronymic
        credentials = "_".join([
            surname,
            name[0] if name else '',
            patronymic[0] if patronymic else '']
        )
        test_name = "_".join(self.test_result.test.name.split())
        file_name = f"{credentials}_{test_name}.docx"
        path = os.path.join(temp_dir, file_name)

        self.doc.save(path)
        return path

    def _set_cell_font_size(self, cell, size):
        """
        Set font size for all runs in cell
        """
        for paragraph in cell.paragraphs:
            self._set_font_size(paragraph, size)

    @staticmethod
    def _set_font_size(paragraph, size):
        """
        Set font size for all runs in paragraph
        """
        for run in paragraph.runs:
            run.font.size = Pt(size)


    def _fill_cell(self, cell, text, align=WD_ALIGN_PARAGRAPH.LEFT):
        """
        Fill cell with text and set alignment
        """
        cell.text = str(text)
        self._set_cell_font_size(cell, 11)
        cell.paragraphs[0].alignment = align
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    def _style_paragraph(self, paragraph):
        """
        Style paragraph with common settings (font size, spacing)
        """
        paragraph.paragraph_format.space_after = 0
        paragraph.paragraph_format.line_spacing = 1.5
        self._set_font_size(paragraph, 12)

    def _add_space_after_table(self):
        """
        Add empty paragraph after table
        """
        paragraph = self.doc.add_paragraph(" ")
        paragraph.paragraph_format.space_after = 0
        paragraph.paragraph_format.space_before = 0
        paragraph.paragraph_format.line_spacing = 0
        self._set_font_size(paragraph, 5)
        return paragraph

    def _create_header(self):
        """
        Create header with information about user and test
        """
        user = self.user_info
        (age, age_ending) = user.get_age()

        header = self.doc.add_paragraph()
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        header.add_run(f"Бланк {self.test_result.test.name}").bold = True
        self._set_font_size(header, 14)

        credentials = self.doc.add_paragraph(f"П.І.Б: {user.surname} "
                                             f"{user.name if user.name else ''} "
                                             f"{user.patronymic if user.patronymic else ''}")
        self._style_paragraph(credentials)

        age = self.doc.add_paragraph(f"Вік пацієнта: {age} {age_ending}")
        self._style_paragraph(age)

        pass_date = self.test_result.passed_at.strftime("%d.%m.%Y %H:%M")
        pass_paragraph = self.doc.add_paragraph(f"Дата проходження: {pass_date}")
        self._style_paragraph(pass_paragraph)